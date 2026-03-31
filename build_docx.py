import os
import re
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def configure_document(doc):
    # Configure page margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Normal style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(14)
    # style_normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = style_normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.all_caps = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(18)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Cm(1.25)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(14)

    # Heading 3
    h3 = doc.styles.add_style('Heading 3 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h3.base_style = h2
    
def clean_md_text(text):
    # simple bold substitution (for inline formatting, we'd need runs, but let's keep it simple for now or implement runs later)
    # Here we won't process inline **bold** perfectly via code blocks if we just want raw string, but let's try to remove ** if necessary
    # Better yet, keep formatting for strong marks if we can.
    return text

def parse_and_append(md_path, doc_path):
    doc = Document(doc_path)
    configure_document(doc)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    started = False
    in_code_block = False
    code_content = []
    
    in_table = False
    table_lines = []

    def flush_code_block():
        if code_content:
            text = "\n".join(code_content).strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(11)
            code_content.clear()

    def flush_table():
        if table_lines:
            # parse markdown table
            # format: | Col1 | Col2 |
            valid_lines = [l for l in table_lines if '|' in l and '---' not in l]
            if not valid_lines:
                table_lines.clear()
                return
            
            rows_data = []
            for line in valid_lines:
                cols = [c.strip() for c in line.strip().strip('|').split('|')]
                rows_data.append(cols)
            
            col_len = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=col_len)
            try:
                table.style = 'Table Grid'
            except KeyError:
                pass # Use default style if Table Grid doesn't exist
            
            for i, r_data in enumerate(rows_data):
                for j, cell_text in enumerate(r_data):
                    table.cell(i, j).text = cell_text.replace('**', '')
                    table.cell(i, j).paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            table_lines.clear()

    # Process lines
    for line in lines:
        line_s = line.strip()
        
        # Start at БОБИ 2
        if line_s.startswith("## БОБИ 2"):
            started = True
        if not started:
            continue

        # Check code block
        if line_s.startswith("```"):
            if in_code_block:
                in_code_block = False
                flush_code_block()
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_content.append(line_s)
            continue
            
        # Check table
        if line_s.startswith("|"):
            in_table = True
            table_lines.append(line_s)
            continue
        else:
            if in_table:
                in_table = False
                flush_table()

        if not line_s:
            continue
            
        # Stop at ХУЛОСА - actually we want to include Chapter 3 and ХУЛОСА and Замимахо if requested?
        # The user said "вторую и последнююю третью главу", meaning we do Chapter 2, 3, Khulosa, Literature etc.
        # It's better to just process til the end.

        text_clean = line.strip()
        
        # Determine paragraph type
        if text_clean.startswith("## "):
            text = text_clean[3:].replace('**', '')
            doc.add_paragraph(text, style='Heading 1')
        elif text_clean.startswith("### "):
            text = text_clean[4:].replace('**', '')
            doc.add_paragraph(text, style='Heading 2')
        elif str(text_clean).startswith("Ҷадвали"):
            # Table caption
            text = text_clean.replace('**', '')
            p = doc.add_paragraph(text)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Cm(0)
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(12)
        elif str(text_clean).startswith("Расми"):
            # Insert prompt placeholder for image first
            img_p = doc.add_paragraph("[Промпт для ИИ: Скриншот для данного рисунка]")
            img_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.paragraph_format.first_line_indent = Cm(0)
            img_p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            # Figure caption below the image
            text = text_clean.replace('**', '')
            p = doc.add_paragraph(text)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(12)
        elif str(text_clean).startswith("Листинги"):
            text = text_clean.replace('**', '')
            p = doc.add_paragraph(text)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(12)
        elif str(text_clean).startswith("ХУЛОСА") or str(text_clean).startswith("ФЕҲРИСТИ АДАБИЁТ") or str(text_clean).startswith("ЗАМИМАҲО"):
            text = text_clean.replace('**', '')
            doc.add_paragraph(text, style='Heading 1')
        else:
            # Inline bold formatting `**` logic
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', text_clean)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
                    
    # Final flushes
    if in_table: flush_table()
    if in_code_block: flush_code_block()

    doc.save("diploma_thesis_updated.docx")
    print("Document formally processed and saved as diploma_thesis_updated.docx")

if __name__ == "__main__":
    import sys
    parse_and_append("/home/sherzod/Рабочий стол/maga/diploma_thesis.md", "/home/sherzod/Рабочий стол/maga/diploma_thesis.docx")
