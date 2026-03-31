import re
from docx import Document
from docx.shared import Inches

def main():
    doc = Document("diploma_thesis.docx")
    
    replace_count = 0
    for i, para in enumerate(doc.paragraphs):
        if "Промпт" in para.text and "для ИИ" in para.text:
            text = para.text.strip()
            para.text = ""
            
            caption = ""
            for j in range(1, 4):
                if i + j < len(doc.paragraphs):
                    if "Расми" in doc.paragraphs[i+j].text:
                        caption = doc.paragraphs[i+j].text
                        break
            
            print(f"Found placeholder. Associated caption snippet: {caption[:40]}")
            replace_count += 1
            
            if "Расми 4" in caption or "data_tajik.csv" in caption:
                para.text = "[Промпт для ИИ: Сгенерируй изображение таблицы данных в формате CSV. Колонки: Source, Relationship, Target, Type, Side. Строки с именами из Шахнаме (Рустам, Сиёвуш, Эрон, Турон).]"
            elif "Расми 5" in caption or "Centrality" in caption:
                try:
                    para.add_run().add_picture("screenshot_tables.png", width=Inches(6.0))
                except Exception as e:
                    para.text = f"[Промпт для ИИ: Скриншот дашборда с двумя таблицами: Degree Centrality и Betweenness Centrality, с метриками.]"
            elif "Расми 6" in caption or "пурраи интерф" in caption:
                try:
                    para.add_run().add_picture("screenshot_full.png", width=Inches(6.0))
                except Exception:
                    para.text = "[Промпт для ИИ: Скриншот полного интерфейса Streamlit с графом Знаний 'Шахнаме'.]"
            elif "Расми 7" in caption or "интерактивии PyVis" in caption:
                try:
                    para.add_run().add_picture("screenshot_pyvis.png", width=Inches(6.0))
                except Exception:
                    para.text = "[Промпт для ИИ: Сгенерируй интерактивный граф связей (зеленые узлы - Эрон, красные - Турон, желтые - локации, фиолетовые - предметы).]"
            elif "Расми 8" in caption or "мизи кор," in caption:
                para.text = "[Промпт для ИИ: Сгенерируй подробную схему (инфографику) эргономичного рабочего места программиста по ГОСТ/СанПиН. Покажи правильное расстояние от глаз до дисплея (60-70 см), высоту стола и ортопедическое кресло.]"
            elif "Расми 9" in caption or "кулли саҳифаи" in caption:
                try:
                    para.add_run().add_picture("screenshot_full.png", width=Inches(6.0))
                except Exception:
                    para.text = "[Промпт для ИИ: Скриншот страницы аналитического дашборда целиком.]"
            elif "Расми 10" in caption or "Танҳо Эрон" in caption:
                try:
                    para.add_run().add_picture("screenshot_pyvis.png", width=Inches(6.0))
                except Exception:
                    para.text = "[Промпт для ИИ: Сгенерируй интерактивный сетевой граф для стороны 'Эрон', где подсвечены в основном ЗЕЛЕНЫЕ узлы.]"
            elif "Расми 11" in caption or "Танҳо Турон" in caption:
                try:
                    para.add_run().add_picture("screenshot_pyvis.png", width=Inches(6.0))
                except Exception:
                    para.text = "[Промпт для ИИ: Сгенерируй интерактивный сетевой граф для стороны 'Турон', где подсвечены в основном КРАСНЫЕ узлы.]"
            else:
                para.text = "[Промпт для ИИ: Диаграмма или изображение по теме текущего раздела]"
                
    doc.save("diploma_thesis_updated.docx")
    print(f"Document saved to diploma_thesis_updated.docx. Replaced {replace_count} placeholders.")

if __name__ == "__main__":
    main()
